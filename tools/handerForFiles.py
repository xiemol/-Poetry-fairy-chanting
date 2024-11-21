from bs4 import BeautifulSoup
import fitz
from langchain_core.documents import Document
import pandas as pd
import os
import csv
from docx import Document as DocxDocument
import chardet


def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read()
    result = chardet.detect(raw_data)
    return result['encoding']


def load_html_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    soup = BeautifulSoup(content, 'html.parser')
    text = soup.get_text(separator='\n')
    return text.strip().replace('\n', ' ')


def create_document_from_html(file_path):
    text = load_html_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_pdf_file(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def create_document_from_pdf(file_path):
    text = load_pdf_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    return text


def create_document_from_txt(file_path):
    text = load_txt_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_csv_file(file_path):
    encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1']  # 常见编码列表
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                reader = csv.reader(file)
                rows = [row for row in reader]
                return rows
        except (UnicodeDecodeError, csv.Error):
            pass
    raise UnicodeDecodeError(f"Could not decode {file_path} using known encodings.")


def create_document_from_csv(file_path):
    rows = load_csv_file(file_path)
    text = "\n".join([",".join(row) for row in rows])
    return [Document(page_content=text, metadata={"source": file_path})]


def load_docx_file(file_path):
    doc = DocxDocument(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def create_document_from_docx(file_path):
    text = load_docx_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def get_all_files(directory):
    files_list = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            files_list.append(os.path.join(root, file))
    return files_list
